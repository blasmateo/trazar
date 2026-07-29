#!/usr/bin/env python3
"""
exportar_docx.py - Herramienta externa de Trazar para exportar
metricas de asistencias (JSON) a documentos .docx.

CONTRATO IPC (Trazar <-> herramientas externas Python):
  - Entrada: envelope JSON por stdin (o --input <RUTA> si pesa).
  - Salida: JSON de resultado por stdout.
  - Logs: texto con prefijos [INFO]/[WARN]/[ERR] por stderr.
  - Artefactos binarios (.docx) al disco, nunca por stdout.
  - Codigo de salida: 0 ok, 1 error de contrato, 2 error de entrada,
    3 error de recursos, 4 error de dependencia.

Envelope de entrada:
  {
    "contractVersion": "1.0",
    "operation": "exportar-docx",
    "payload": { "datos": <JSON de metricas> },
    "output": { "ruta": "<absoluta>", "modo": "lista|tabla" }
  }

Respuesta de salida (stdout):
  { "status": "ok", "artefactos": [{"ruta": "...", "tipo": "docx", "tamanoBytes": N}] }
  o { "status": "error", "error": {"codigo": "...", "mensaje": "..."} }

Compatibilidad legacy: si se invoca con --json/--salida/--modo (flags
sueltas), se acepta por compatibilidad con flujos que no usen el envelope.
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

CONTRACT_VERSION = "1.0"


def log_info(msg):
    print(f"[INFO] {msg}", file=sys.stderr)


def log_warn(msg):
    print(f"[WARN] {msg}", file=sys.stderr)


def log_err(msg):
    print(f"[ERR] {msg}", file=sys.stderr)


_CODIGOS_SALIDA = {
    "E_CONTRATO": 1,
    "E_ENTRADA": 2,
    "E_RECURSOS": 3,
    "E_DEPENDENCIA": 4,
}


def responder(status, artefactos=None, error=None):
    """Escribe el JSON de resultado por stdout y termina con codigo semantico."""
    out = {"contractVersion": CONTRACT_VERSION, "status": status}
    if artefactos:
        out["artefactos"] = artefactos
    if error:
        out["error"] = error
    print(json.dumps(out, ensure_ascii=False))
    if status == "ok":
        sys.exit(0)
    cod = error.get("codigo", "E_DESCONOCIDO") if error else "E_DESCONOCIDO"
    sys.exit(_CODIGOS_SALIDA.get(cod, 1))


def main():
    args = sys.argv[1:]

    # Deteccion del modo: flags legacy (--json) vs contrato (envelope).
    if "--json" in args or any(a.startswith("--json=") for a in args):
        _modo_legacy()
        return

    # --- Modo contrato: leer envelope desde stdin o --input ---
    ruta_input = None
    for i, a in enumerate(args):
        if a.startswith("--input="):
            ruta_input = a.split("=", 1)[1]
        elif a == "--input" and i + 1 < len(args):
            ruta_input = args[i + 1]

    try:
        if ruta_input:
            log_info(f"Leyendo envelope desde archivo: {ruta_input}")
            with open(ruta_input, "r", encoding="utf-8") as f:
                envelope = json.load(f)
        else:
            log_info("Leyendo envelope desde stdin")
            envelope = json.load(sys.stdin)
    except json.JSONDecodeError as e:
        log_err(f"JSON de entrada invalido: {e}")
        responder("error", error={"codigo": "E_CONTRATO", "mensaje": f"JSON invalido: {e}"})
        return
    except OSError as e:
        log_err(f"No se pudo leer la entrada: {e}")
        responder("error", error={"codigo": "E_ENTRADA", "mensaje": str(e)})
        return


    # Validar contrato
    version = envelope.get("contractVersion")
    if version != CONTRACT_VERSION:
        log_err(f"Version de contrato no soportada: {version} (esperada {CONTRACT_VERSION})")
        responder("error", error={"codigo": "E_CONTRATO",
                                  "mensaje": f"Version de contrato {version} no soportada (esperada {CONTRACT_VERSION})"})
        return

    payload = envelope.get("payload", {})
    output = envelope.get("output", {})
    datos = payload.get("datos")
    ruta_salida = output.get("ruta")
    modo = output.get("modo", "lista")

    if datos is None:
        log_err("El payload no contiene 'datos'")
        responder("error", error={"codigo": "E_CONTRATO", "mensaje": "Falta 'payload.datos'"})
        return
    if not ruta_salida:
        log_err("El output no contiene 'ruta'")
        responder("error", error={"codigo": "E_CONTRATO", "mensaje": "Falta 'output.ruta'"})
        return
    if modo not in ("lista", "tabla"):
        log_err(f"Modo invalido: {modo}")
        responder("error", error={"codigo": "E_CONTRATO", "mensaje": f"Modo '{modo}' invalido (lista|tabla)"})
        return

    # La ruta de salida debe ser un archivo, no un directorio.
    salida = Path(ruta_salida)
    if salida.is_dir():
        log_err(f"La ruta de salida es un directorio: {salida}")
        responder("error", error={"codigo": "E_ENTRADA",
                                  "mensaje": (f"La ruta de salida es un directorio, no un archivo: {salida}\n"
                                              f"  Indique el nombre del documento, por ejemplo: {salida / 'asistencias.docx'}")})
        return

    log_info(f"Generando documento {modo} en {salida}")
    try:
        if modo == "lista":
            exportar_lista(datos, str(salida))
        else:
            exportar_tabla(datos, str(salida))
    except Exception as e:
        log_err(f"No se pudo crear el documento: {e}")
        responder("error", error={"codigo": "E_RECURSOS",
                                  "mensaje": (f"No se pudo crear el documento: {e}\n"
                                              f"  Verifique que la ruta sea un .docx valido y tenga permisos de escritura.")})
        return

    tamano = salida.stat().st_size if salida.exists() else 0
    log_info(f"Documento generado ({tamano} bytes)")
    responder("ok", artefactos=[{"ruta": str(salida), "tipo": "docx", "tamanoBytes": tamano}])


def _modo_legacy():
    """Modo de flags sueltas por compatibilidad (--json/--salida/--modo)."""
    import argparse

    parser = argparse.ArgumentParser(description="Exportar metricas de asistencias a .docx (legacy)")
    parser.add_argument("--json", required=True, help="Ruta del archivo JSON de metricas")
    parser.add_argument("--salida", required=True, help="Ruta de salida del archivo .docx")
    parser.add_argument("--modo", required=True, choices=["lista", "tabla"])
    args = parser.parse_args()

    ruta_json = Path(args.json)
    if not ruta_json.exists():
        log_err(f"No existe el archivo JSON: {ruta_json}")
        sys.exit(2)

    try:
        with open(ruta_json, "r", encoding="utf-8-sig") as f:
            datos = json.load(f)
    except json.JSONDecodeError as e:
        log_err(f"Error al parsear JSON: {e}")
        sys.exit(2)

    salida = Path(args.salida)
    if salida.is_dir():
        log_err(f"La ruta de salida es un directorio: {salida}\n"
                f"  Indique el nombre del documento, por ejemplo: {salida / 'asistencias.docx'}")
        sys.exit(2)

    log_info(f"Generando documento {args.modo} en {salida}")
    try:
        if args.modo == "lista":
            exportar_lista(datos, args.salida)
        else:
            exportar_tabla(datos, args.salida)
    except Exception as e:
        log_err(f"No se pudo crear el documento: {e}")
        sys.exit(3)

    print(f"\u2713 Documento exportado: {args.salida}")



def _cabecera(doc, datos):
    """Escribe la cabecera comun del documento."""
    curso = datos.get("curso", "Curso")
    total_clases = datos.get("totalClases", 0)
    fecha = datos.get("fechaActualizacion", "")

    titulo = doc.add_heading(f"Metricas de Asistencias - {curso}", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("Total de clases: ").bold = True
    p.add_run(f"{total_clases}")

    if fecha:
        p2 = doc.add_paragraph()
        p2.add_run("Fecha de actualizacion: ").bold = True
        p2.add_run(fecha)

    doc.add_paragraph()


def exportar_lista(datos, salida):
    """Exporta el modo lista (resumen por cursante, sin detalle)."""
    doc = Document()
    _cabecera(doc, datos)

    cursantes = datos.get("cursantes", {})
    if not cursantes:
        doc.add_paragraph("No hay cursantes para mostrar.")
    else:
        doc.add_heading(f"Cursantes ({len(cursantes)})", level=2)

        tabla = doc.add_table(rows=1, cols=4)
        tabla.style = "Light Grid Accent 1"
        hdr = tabla.rows[0].cells
        hdr[0].text = "Cursante"
        hdr[1].text = "Presente"
        hdr[2].text = "Ausente"
        hdr[3].text = "% Asistencia"

        for nombre in sorted(cursantes.keys()):
            info = cursantes[nombre]
            presente = int(info.get("presente", 0))
            ausente = int(info.get("ausente", 0))
            total = presente + ausente
            porcentaje = round((presente / total * 100)) if total > 0 else 0

            fila = tabla.add_row().cells
            fila[0].text = nombre
            fila[1].text = str(presente)
            fila[2].text = str(ausente)
            fila[3].text = f"{porcentaje}%"

    doc.save(salida)


def exportar_tabla(datos, salida):
    """Exporta el modo tabla (detalle por cursante con asignatura/clase/asiste)."""
    doc = Document()
    _cabecera(doc, datos)

    cursantes = datos.get("cursantes", {})
    if not cursantes:
        doc.add_paragraph("No hay cursantes para mostrar.")
    else:
        for nombre in sorted(cursantes.keys()):
            info = cursantes[nombre]
            presente = int(info.get("presente", 0))
            ausente = int(info.get("ausente", 0))
            total = presente + ausente
            porcentaje = round((presente / total * 100)) if total > 0 else 0

            doc.add_heading(f"{nombre}", level=2)
            p = doc.add_paragraph()
            p.add_run(
                f"Total: {total} | Presente: {presente} | "
                f"Ausente: {ausente} | % Asistencia: {porcentaje}%"
            )

            detalle = info.get("detalle", [])
            if detalle:
                tabla = doc.add_table(rows=1, cols=3)
                tabla.style = "Light Grid Accent 1"
                hdr = tabla.rows[0].cells
                hdr[0].text = "Asignatura"
                hdr[1].text = "Clase"
                hdr[2].text = "Asiste"

                def ordenar(asis):
                    asig = asis.get("asignatura", "")
                    id_str = asig.split("-")[0] if "-" in asig else asig
                    try:
                        id_num = int(id_str)
                    except ValueError:
                        id_num = 999999
                    return (id_num, int(asis.get("clase", 0)))

                for asis in sorted(detalle, key=ordenar):
                    asig = asis.get("asignatura", "")
                    asig_mostrar = asig if asig else "eventos-academicos"
                    clase = int(asis.get("clase", 0))
                    presente = bool(asis.get("presente", False))
                    estado = "si" if presente else "no"

                    fila = tabla.add_row().cells
                    fila[0].text = asig_mostrar
                    fila[1].text = str(clase)
                    fila[2].text = estado

            doc.add_paragraph()  # separacion entre cursantes

    doc.save(salida)


if __name__ == "__main__":
    main()
